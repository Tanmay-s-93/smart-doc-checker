from flask import Flask, request, render_template, jsonify
import os
from pypdf import PdfReader
from docx import Document
from dotenv import load_dotenv
import google.generativeai as genai
from werkzeug.utils import secure_filename
import magic  # python-magic library for file type validation
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_cors import CORS
import logging
from logging.handlers import RotatingFileHandler
import time

# Load environment variables
load_dotenv()

# Initialize Flask app
app = Flask(__name__)

# --- Configuration ---
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'docx'}
app.config['RATE_LIMIT_STORAGE_URI'] = 'memory://'

# --- CORS Support ---
CORS(app)  # Enable CORS for all routes

# --- Rate Limiting ---
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["200 per day", "50 per hour"],
    storage_uri=app.config['RATE_LIMIT_STORAGE_URI']
)

# --- Logging Configuration ---
if not os.path.exists('logs'):
    os.makedirs('logs')

file_handler = RotatingFileHandler('logs/app.log', maxBytes=10240, backupCount=10)
file_handler.setFormatter(logging.Formatter(
    '%(asctime)s %(levelname)s: %(message)s [in %(pathname)s:%(lineno)d]'
))
file_handler.setLevel(logging.INFO)
app.logger.addHandler(file_handler)
app.logger.setLevel(logging.INFO)
app.logger.info('Smart Doc Checker startup')

# --- API Key Validation ---
api_key = os.getenv("GOOGLE_API_KEY")
if not api_key:
    app.logger.warning("GOOGLE_API_KEY not found in environment variables")
else:
    genai.configure(api_key=api_key)

# In-memory "database" for the usage counter demo
usage_data = {
    'reports_generated': 0,
    'last_reset': time.time()
}

# --- Helper Functions ---
def allowed_file(filename):
    """Check if the file has an allowed extension"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def validate_file_type(file_stream, filename):
    """Validate file type using magic numbers"""
    try:
        # Read the first 2048 bytes for MIME type detection
        header = file_stream.read(2048)
        file_stream.seek(0)  # Reset stream position
        
        mime_type = magic.from_buffer(header, mime=True)
        
        if filename.endswith('.pdf'):
            return 'pdf' in mime_type
        elif filename.endswith('.docx'):
            return 'officedocument' in mime_type or 'word' in mime_type.lower()
        
        return False
    except Exception as e:
        app.logger.error(f"Error validating file type: {str(e)}")
        return False

def extract_text_from_file(file):
    """Extract text from uploaded PDF or DOCX files with enhanced error handling"""
    if not file or not file.filename:
        return ""
    
    text = ""
    filename = secure_filename(file.filename.lower())
    
    try:
        if filename.endswith('.pdf'):
            try:
                pdf_reader = PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    else:
                        # OCR fallback message
                        text += "[Could not extract text from this page - may be a scanned image]\n"
            except Exception as pdf_error:
                app.logger.error(f"PDF extraction error: {str(pdf_error)}")
                # Try alternative extraction method
                try:
                    # Alternative PDF extraction could be implemented here
                    raise Exception(f"PDF extraction failed: {str(pdf_error)}")
                except Exception as alt_error:
                    raise Exception(f"Could not extract text from PDF: {str(alt_error)}")
                    
        elif filename.endswith('.docx'):
            try:
                doc = Document(file)
                for para in doc.paragraphs:
                    text += para.text + "\n"
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + "\t"
                        text += "\n"
                    text += "\n"
            except Exception as docx_error:
                app.logger.error(f"DOCX extraction error: {str(docx_error)}")
                raise Exception(f"Could not extract text from Word document: {str(docx_error)}")
        else:
            raise ValueError("Unsupported file format.")
            
    except Exception as e:
        app.logger.error(f"Error reading {filename}: {str(e)}")
        raise Exception(f"Error reading {filename}: {str(e)}")
    
    if not text.strip():
        raise Exception(f"No extractable text found in {filename}. The document might be scanned, encrypted, or contain only images.")
    
    return text

def find_contradictions(text1, text2, text3=""):
    """Use the Google Gemini API to analyze texts for contradictions."""
    # Check if API key is available
    if not os.getenv("GOOGLE_API_KEY"):
        return "Error: Google API key not configured. Please check your environment variables."
    
    combined_text = f"DOCUMENT 1:\n{text1[:3000]}\n\nDOCUMENT 2:\n{text2[:3000]}"
    
    if text3:
        combined_text += f"\n\nEXTERNAL POLICY UPDATE (DOCUMENT 3):\n{text3[:3000]}"

    prompt = f"""
    Act as a meticulous legal and policy analyst. Compare the following documents and identify any clear contradictions, conflicts, or inconsistencies.

    {combined_text}

    Provide a structured report. For each contradiction found:
    1. Quote the conflicting sections from each document.
    2. Clearly explain the nature of the contradiction.
    3. Suggest a clarification to resolve the conflict.

    If no contradictions are found, clearly state "No contradictions found."
    """
    
    try:
        model = genai.GenerativeModel('gemini-1.5-flash-latest')
        response = model.generate_content(prompt)
        
        # More robust error checking for Gemini
        if not response.parts:
            try:
                # Check for specific safety feedback if available
                if hasattr(response, 'prompt_feedback') and response.prompt_feedback.block_reason:
                    error_msg = f"Analysis blocked. Reason: {response.prompt_feedback.block_reason.name}"
                    app.logger.warning(error_msg)
                    return error_msg
            except (AttributeError, ValueError):
                pass  # Fallback to generic message
            error_msg = "The AI model returned an empty response. This may be due to a safety filter or the content of the documents."
            app.logger.warning(error_msg)
            return error_msg
            
        return response.text
        
    except Exception as e:
        error_msg = f"An error occurred with the AI analysis: {str(e)}"
        app.logger.error(error_msg)
        return error_msg

# --- Routes ---
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/health')
def health_check():
    """Health check endpoint for monitoring"""
    api_status = "configured" if os.getenv("GOOGLE_API_KEY") else "not configured"
    return jsonify({
        'status': 'healthy', 
        'service': 'Smart Doc Checker API',
        'api_key_status': api_status,
        'reports_generated': usage_data['reports_generated'],
        'uptime': round(time.time() - usage_data['last_reset'], 2)
    })

@app.route('/analyze', methods=['POST'])
@limiter.limit("10 per minute")  # Rate limiting for analysis endpoint
def analyze_documents():
    # Check for API key before processing
    if not os.getenv("GOOGLE_API_KEY"):
        return jsonify({'error': 'Google API key not configured. Please contact administrator.'}), 500
        
    file1 = request.files.get('file1')
    file2 = request.files.get('file2')
    file3 = request.files.get('file3')

    if not file1 or not file2 or file1.filename == '' or file2.filename == '':
        return jsonify({'error': 'Please upload at least two documents.'}), 400

    # File Validation Logic
    files_to_check = [(file1, 'file1'), (file2, 'file2')]
    if file3 and file3.filename != '':
        files_to_check.append((file3, 'file3'))
    
    for file, field_name in files_to_check:
        if file and file.filename:
            filename = secure_filename(file.filename)
            
            # Check extension
            if not allowed_file(filename):
                return jsonify({'error': f"Invalid file extension for {filename}. Only PDF and DOCX are allowed."}), 400
            
            # Check MIME type to prevent spoofing
            if not validate_file_type(file.stream, filename):
                return jsonify({'error': f"Invalid file content for {filename}. The file is not a valid PDF or DOCX."}), 400
            
            # Reset stream position after validation
            file.stream.seek(0)
    
    try:
        text1 = extract_text_from_file(file1)
        text2 = extract_text_from_file(file2)
        text3 = extract_text_from_file(file3) if file3 and file3.filename != '' else ""

        # Check if we extracted any text
        if not text1.strip() or not text2.strip():
            return jsonify({'error': 'Could not extract text from one or more documents. They might be scanned images or encrypted.'}), 400

        report = find_contradictions(text1, text2, text3)
        
        usage_data['reports_generated'] += 1
        
        return jsonify({
            'report': report, 
            'usage_count': usage_data['reports_generated'],
            'document1_length': len(text1),
            'document2_length': len(text2),
            'document3_length': len(text3) if text3 else 0
        })

    except Exception as e:
        app.logger.error(f"Analysis error: {str(e)}")
        return jsonify({'error': str(e)}), 500

# --- Error Handlers ---
@app.errorhandler(413)
def too_large(e):
    app.logger.warning(f"File too large: {str(e)}")
    return jsonify({'error': 'File too large. Maximum size is 16MB.'}), 413

@app.errorhandler(429)
def ratelimit_handler(e):
    app.logger.warning(f"Rate limit exceeded: {str(e)}")
    return jsonify({'error': 'Rate limit exceeded. Please try again later.'}), 429

@app.errorhandler(500)
def internal_error(e):
    app.logger.error(f"Internal server error: {str(e)}")
    return jsonify({'error': 'Internal server error. Please try again later.'}), 500

if __name__ == '__main__':
    # Create templates directory if it doesn't exist
    if not os.path.exists('templates'):
        os.makedirs('templates')
    
    app.logger.info("Starting Smart Doc Checker server")
    app.run(debug=True, host='0.0.0.0', port=5000)